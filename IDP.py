# ==============================
# INTELLIGENT DOCUMENT PROCESSOR - with Metrics and Logo
# ==============================

import os
import base64
import tempfile
import json
import re
from pathlib import Path
from io import BytesIO

import streamlit as st
import pandas as pd

import os
import streamlit as st

#os.environ["OPENAI_API_KEY"] = st.secrets["OPENAI_API_KEY"]

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_openai import OpenAIEmbeddings, ChatOpenAI

from langchain_community.document_loaders import (
    TextLoader,
    PyPDFLoader,
    UnstructuredPowerPointLoader,
    UnstructuredExcelLoader,
)

from langchain_core.documents import Document
from langchain_core.messages import HumanMessage
from docx import Document as DocxDocument
from streamlit_pdf_viewer import pdf_viewer

# Use API key through input from user
def get_llm():
    return ChatOpenAI(
        model="gpt-4o",
        temperature=0,
        api_key=st.session_state["api_key"]
    )

def get_embeddings():
    return OpenAIEmbeddings(
        api_key=st.session_state["api_key"]
    )

# ------------------------------
# INIT
# ------------------------------

st.set_page_config("IDP - Professional", layout="wide")

# ------------------------------
# LOGIN + API KEY VALIDATION
# ------------------------------

import streamlit as st
from pathlib import Path
from openai import OpenAI

USERS = st.secrets.get("users", {})

# ------------------------------
# VALIDATE API KEY
# ------------------------------
def validate_api_key(api_key):
    try:
        client = OpenAI(api_key=api_key)

        # Lightweight test call
        client.models.list()

        return True
    except Exception:
        return False


# ------------------------------
# LOGIN FUNCTION
# ------------------------------
def login():
    logo_path = Path(__file__).parent / "IDP-Logo1.png"

    col1, col2, col3 = st.columns([1, 2, 1])

    with col2:
        if logo_path.exists():
            st.image(logo_path, width=220)

        st.markdown("### 🔐 Sign In")

        username = st.text_input("Username")
        password = st.text_input("Password", type="password")
        api_key = st.text_input("OpenAI API Key", type="password")

        if st.button("Login", use_container_width=True):

            # Validate user
            if username not in USERS or USERS[username]["password"] != password:
                st.error("Invalid username or password")
                return

            if not api_key:
                st.error("Please enter your OpenAI API key")
                return

            # 🔑 Validate API key
            with st.spinner("Validating API key..."):
                if not validate_api_key(api_key):
                    st.error("Invalid OpenAI API key")
                    return

            # Save session
            st.session_state["logged_in"] = True
            st.session_state["user"] = username
            st.session_state["role"] = USERS[username].get("role", "user")
            st.session_state["api_key"] = api_key

            st.success(f"Welcome {username}")
            st.rerun()

# ------------------------------
# SESSION INIT
# ------------------------------
if "logged_in" not in st.session_state:
    st.session_state["logged_in"] = False

if "user" not in st.session_state:
    st.session_state["user"] = None

if "role" not in st.session_state:
    st.session_state["role"] = None

if "api_key" not in st.session_state:
    st.session_state["api_key"] = None


# ------------------------------
# LOGIN GATE
# ------------------------------
if not st.session_state["logged_in"]:
    login()
    st.stop()


# ------------------------------
# SIDEBAR (USER INFO + LOGOUT)
# ------------------------------
with st.sidebar:
    st.markdown("### 👤 User Info")
    st.write(f"**User:** {st.session_state['user']}")
    st.write(f"**Role:** {st.session_state['role']}")

    st.success("🔑 API key loaded securely")

    if st.button("🚪 Logout"):
        for key in ["logged_in", "user", "role", "api_key"]:
            if key in st.session_state:
                del st.session_state[key]

        st.success("Logged out")
        st.rerun()

logo_path = Path(__file__).parent / "IDP-Logo1.png"

col1, col2 = st.columns([1, 7], gap="small")

with col1:
    st.image(logo_path, width=280)

with col2:
    st.markdown("## Intelligent Document Processor")
    st.caption("AI-powered document understanding & automation")

# Session state
for key in ["structured_data", "doc_type", "vectorstore", "full_text"]:
    if key not in st.session_state:
        st.session_state[key] = None

if "generated_resume" not in st.session_state:
    st.session_state.generated_resume = None

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

if "processed_file" not in st.session_state:
    st.session_state.processed_file = None

if "active_tab" not in st.session_state:
    st.session_state.active_tab = 0

if "suggested_questions" not in st.session_state:
    st.session_state.suggested_questions = []

if "metrics" not in st.session_state:
    st.session_state.metrics = {
        "tokens": 0,
        "response_times": [],
        "accuracy_scores": []
    }

if "doc_metrics" not in st.session_state:
    st.session_state.doc_metrics = {}

# ------------------------------
# FILE UPLOAD
# ------------------------------

uploaded_file = st.file_uploader(
    "Drag and drop file here",
    type=["txt", "pdf", "docx", "pptx", "xlsx", "png", "jpg", "jpeg"]
)

# ------------------------------
# HELPERS
# ------------------------------

def save_temp_file(uploaded_file):
    suffix = Path(uploaded_file.name).suffix
    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(uploaded_file.getvalue())
        return tmp.name


def load_docx_safe(file_path):
    doc = DocxDocument(file_path)
    text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    return [Document(page_content=text)]


def process_file(uploaded_file):
    documents = []
    if not uploaded_file:
        return documents

    suffix = Path(uploaded_file.name).suffix.lower()

    if suffix in [".png", ".jpg", ".jpeg"]:
        encoded = base64.b64encode(uploaded_file.getvalue()).decode()

        message = HumanMessage(content=[
            {"type": "text", "text": "Extract all readable text with structure (headings, tables, key-value pairs)."},
            {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{encoded}"}}
        ])

        response = llm.invoke([message])
        documents.append(Document(page_content=response.content))

    else:
        file_path = save_temp_file(uploaded_file)

        if suffix == ".txt":
            try:
                documents.extend(TextLoader(file_path, encoding="utf-8").load())
            except Exception:
                documents.extend(TextLoader(file_path, encoding="cp1252").load())
        elif suffix == ".pdf":
            documents.extend(PyPDFLoader(file_path).load())
        elif suffix == ".docx":
            documents.extend(load_docx_safe(file_path))
        elif suffix == ".pptx":
            documents.extend(UnstructuredPowerPointLoader(file_path).load())
        elif suffix == ".xlsx":
            documents.extend(UnstructuredExcelLoader(file_path).load())

    return documents


def safe_json_parse(response):
    try:
        return json.loads(response)
    except:
        match = re.search(r"\{.*\}", response, re.DOTALL)
        if match:
            try:
                return json.loads(match.group())
            except:
                pass
    return {"error": "Invalid JSON output", "raw_response": response[:500]}


def detect_document_type(text):
    prompt = f"""
Classify document into ONE:
Resume, Invoice, Receipt, Report, Ticket, Other
Return only the label.
{text[:2000]}
"""
    raw = tracked_llm_call(prompt).content.strip().lower()

    if raw.startswith("resume"):
        return "resume"
    elif raw.startswith("invoice"):
        return "invoice"
    elif raw.startswith("receipt"):
        return "receipt"
    elif raw.startswith("report"):
        return "report"
    elif raw.startswith("ticket"):
        return "ticket"
    else:
        return "other"


def extract_structured_json(text, doc_type):

    # ✅ Clean text
        clean_text = re.sub(r"[^\x00-\x7F]+", " ", text)
        clean_text = clean_text.replace("{", "").replace("}", "")

        # 🎯 Different prompt based on document type
        if doc_type == "resume":

            prompt = f"""
        You are a strict JSON generator.

        Return ONLY valid JSON.

        MANDATORY SCHEMA:
        {{
        "name": "",
        "email": "",
        "phone": "",
        "skills": [],
        "education": [],
        "experience": []
        }}

        RULES:
        - ALWAYS include "name"
        - Do not add extra keys

        Content:
        {clean_text[:4000]}
        """

        else:
            # ✅ Invoice / Report / Ticket → FULL EXTRACTION
            prompt = f"""
        Extract ALL possible key-value pairs from the document.

        Return ONLY valid JSON.

        RULES:
        - Capture every identifiable field
        - Preserve original field names where possible
        - Include nested structures if present
        - Do NOT summarize
        - Do NOT skip fields

        Examples:
        Invoice → invoice_number, date, total, vendor, line_items
        Report → title, author, summary, sections
        Ticket → issue, status, priority, requester

        Content:
        {clean_text[:4000]}
        """

        try:
            response = tracked_llm_call(prompt).content.replace("```json", "").replace("```", "").strip()
            parsed = safe_json_parse(response)

            # ✅ Ensure parsed is always a dictionary
            if isinstance(parsed, list):
                try:
                    merged = {}
                    for item in parsed:
                        if isinstance(item, dict):
                            merged.update(item)
                    parsed = merged
                except:
                    parsed = {"data": parsed}

            # AI fallback if name missing
            if doc_type == "resume" and isinstance(parsed, dict) and not parsed.get("name"):
                try:
                    fallback_prompt = f"""
            Extract only the full name of the person from this text.
            Return only the name. No explanation.

            {text[:2000]}
            """
                    fallback_name = tracked_llm_call(fallback_prompt).content.strip()
                    parsed["name"] = fallback_name
                except:
                    parsed["name"] = "candidate"

            # ✅ Optional: format name nicely
            if parsed.get("name"):
                parsed["name"] = parsed["name"].title()

            return parsed


        except Exception as e:
            return {
                "error": "LLM request failed",
                "details": str(e)[:300]
            }


def json_to_kv_dataframe(data):
    rows = []

    def flatten(prefix, obj):
        if isinstance(obj, dict):
            for k, v in obj.items():
                flatten(f"{prefix}.{k}" if prefix else k, v)
        elif isinstance(obj, list):
            for i, item in enumerate(obj):
                flatten(f"{prefix}[{i}]", item)
        else:
            rows.append({"Field": prefix, "Value": json.dumps(obj) if isinstance(obj, (dict, list)) else str(obj)})

    flatten("", data)
    return pd.DataFrame(rows)


def generate_excel(df):
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='data')
    return output.getvalue()

# Resume helpers

def replace_placeholders(doc, placeholders):
    for para in doc.paragraphs:
        for run in para.runs:
            for key, value in placeholders.items():
                if key in run.text:
                    run.text = run.text.replace(key, str(value))


def generate_resume_summary(data):
    prompt = f"""
Create a professional resume summary.
Write candidate name at the top.
Write education, certification and expereince in concise bullet points.
STRICT RULES:
- No markdown
- No ** or *
- Plain text only
{json.dumps(data)}
"""
    return tracked_llm_call(prompt).content


def build_resume(data, template_file):
    summary = generate_resume_summary(data)

    if template_file:
        path = save_temp_file(template_file)
        doc = DocxDocument(path)
    else:
        doc = DocxDocument()

    placeholders = {
        "{{name}}": data.get("name", ""),
        "{{email}}": data.get("email", ""),
        "{{phone}}": data.get("phone", ""),
        "{{summary}}": summary,
    }

    replace_placeholders(doc, placeholders)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

import uuid

def create_vectorstore(docs):
    splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=150)
    chunks = splitter.split_documents(docs)
    db_path = f"./chroma_db_{uuid.uuid4().hex}"
    return Chroma.from_documents(chunks, embedding=embeddings, persist_directory=db_path)

import time

def tracked_llm_call(prompt):
    start = time.time()

    response = get_llm().invoke(prompt)

    duration = time.time() - start
    tokens = len(str(prompt)) // 4 + len(str(response.content)) // 4

    # Global metrics
    st.session_state.metrics["tokens"] += tokens
    st.session_state.metrics["response_times"].append(duration)

    # ✅ Per-document metrics
    current_file = st.session_state.get("current_file")

    if current_file and current_file in st.session_state.doc_metrics:
        doc_metric = st.session_state.doc_metrics[current_file]
        doc_metric["tokens"] += tokens
        doc_metric["response_times"].append(duration)
        doc_metric["calls"] += 1

    return response

# ------------------------------
# PROCESSING WITH PROGRESS
# ------------------------------

if uploaded_file:

    if st.session_state.get("processed_file") != uploaded_file.name:

        current_file = uploaded_file.name
        st.session_state.current_file = current_file

        if current_file not in st.session_state.doc_metrics:
            st.session_state.doc_metrics[current_file] = {
                "tokens": 0,
                "response_times": [],
                "calls": 0
            }
       
        progress = st.progress(0, text="Processing Started...")

        docs = process_file(uploaded_file)
        progress.progress(20, text="File processed")

        st.session_state.full_text = "\n".join([d.page_content for d in docs])
        progress.progress(40, text="Text extracted")

        st.session_state.doc_type = detect_document_type(st.session_state.full_text)
        progress.progress(60, text="Document type detected")

        st.session_state.structured_data = extract_structured_json(
            st.session_state.full_text,
            st.session_state.doc_type
        )
        progress.progress(80, text="Structured data extracted")

        st.session_state.vectorstore = create_vectorstore(docs)
        progress.progress(100, text="Vector index created")

        # 🎯 Auto-suggested questions
        doc_type = st.session_state.doc_type

        if doc_type == "invoice":
            st.session_state.suggested_questions = [
                "What is the total amount?",
                "Who is the vendor?",
                "What is the invoice date?",
                "List all line items"
            ]

        elif doc_type == "resume":
            st.session_state.suggested_questions = [
                "Summarize this candidate",
                "What skills does the candidate have?",
                "What is the experience?",
                "What is the education background?"
            ]

        elif doc_type == "report":
            st.session_state.suggested_questions = [
                "Summarize this report",
                "What are the key findings?",
                "Who is the author?",
                "What are the main sections?"
            ]

        else:
            st.session_state.suggested_questions = [
                "Summarize this document",
                "What are the key points?",
                "Extract important information"
            ]

        st.session_state.processed_file = uploaded_file.name
        #progress.empty()

    st.success(f"✅ Processed Successfully | Type: {st.session_state.doc_type.upper()}")


# ------------------------------
# TABS
# ------------------------------

#tabs = ["Preview", "JSON", "Chat", "Download", "Concur"]
tabs = ["Preview", "JSON", "Chat", "Download", "Concur", "Metrics"]

selected_tab = st.radio(
    "",
    tabs,
    horizontal=True,
    key="active_tab"
)

# PREVIEW
if selected_tab == "Preview":
    if uploaded_file:
        if "pdf" in uploaded_file.type:
            pdf_viewer(uploaded_file.getvalue(), height=200)
        elif "image" in uploaded_file.type:
            st.image(uploaded_file, width=300)
        elif "text" in uploaded_file.type:
            try:
                preview_text = uploaded_file.getvalue().decode("utf-8")
            except:
                preview_text = uploaded_file.getvalue().decode("cp1252", errors="ignore")
            st.text_area("Preview", preview_text, height=200)
        elif "word" in uploaded_file.type or uploaded_file.name.endswith(".docx"):
            path = save_temp_file(uploaded_file)
            doc = DocxDocument(path)
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            st.text_area("DOCX Preview", text, height=200)

# JSON
if selected_tab == "JSON":
    if st.session_state.structured_data:
        st.json(st.session_state.structured_data)

# CHAT
if selected_tab == "Chat":
    if st.session_state.vectorstore:
        for msg in st.session_state.chat_history:
            with st.chat_message(msg["role"]):
                st.write(msg["content"])

        # 🎯 Suggested Questions UI
        if st.session_state.suggested_questions:
            st.markdown("### 💡 Suggested Questions")

            cols = st.columns(len(st.session_state.suggested_questions))

            for i, q in enumerate(st.session_state.suggested_questions):
                if cols[i].button(q):
                    st.session_state.chat_history.append({"role": "user", "content": q})

                    docs = st.session_state.vectorstore.similarity_search(q, k=3)
                    context = "\n\n".join([d.page_content for d in docs])

                    response = tracked_llm_call(
                        f"Answer strictly from context.\nContext:\n{context}\nQ:{q}"
                    ).content

                    st.session_state.chat_history.append({"role": "assistant", "content": response})
                    st.write(response)


        query = st.chat_input("Ask a question")

        if query:
            st.session_state.chat_history.append({"role": "user", "content": query})
            docs = st.session_state.vectorstore.similarity_search(query, k=3)
            context = "\n\n".join([d.page_content for d in docs])
            response = tracked_llm_call(f"Answer strictly from context.\nContext:\n{context}\nQ:{query}").content
            st.session_state.chat_history.append({"role": "assistant", "content": response})
            st.write(response)
            

# DOWNLOAD
if selected_tab == "Download":
    if st.session_state.structured_data:
        st.download_button("Download JSON", json.dumps(st.session_state.structured_data, indent=2), "data.json")

        if st.session_state.doc_type == "invoice":

            df = json_to_kv_dataframe(st.session_state.structured_data)
            st.dataframe(df)

            excel = generate_excel(df)

            # 🎯 Extract meaningful invoice name
            data = st.session_state.structured_data

            invoice_name = (
                data.get("invoice_number")
                or data.get("invoice_no")
                or data.get("invoice_id")
                or data.get("bill_number")
                or data.get("vendor")
                or data.get("supplier")
                or data.get("name")
                or "invoice_data"
            )

            # Clean filename
            safe_name = re.sub(r'[\\/*?:"<>|]', "", str(invoice_name))

            file_name = f"{safe_name}.xlsx"

            # ✅ Show filename
            st.caption(f"📄 {file_name}")

            st.download_button(
                "Download Excel",
                excel,
                file_name
            )

        if st.session_state.doc_type == "resume":

            template_file = st.file_uploader("Upload Resume Template", type=["docx"])

            if template_file:
                st.session_state.generated_resume = build_resume(
                    st.session_state.structured_data,
                    template_file
                )

            if st.session_state.generated_resume:

                data = st.session_state.structured_data

                name = (
                    data.get("name")
                    or data.get("Name")
                    or data.get("candidate_name")
                    or (data.get("personal_details", {}).get("name") if isinstance(data.get("personal_details"), dict) else None)
                    or "candidate"
                )

                safe_name = re.sub(r'[\\/*?:"<>|]', "", name)

            # ✅ ADD THIS to display created file name
                file_name = f"{safe_name}.docx"
                st.caption(f"📄 {file_name}")

                st.download_button(
                    "Download Resume",
                    st.session_state.generated_resume,
                    f"{safe_name}.docx"
                )

# CONCUR
if selected_tab == "Concur":
    st.subheader("Send to Concur Integration")

    supported_types = ["invoice", "ticket"]
    mode = st.radio("Mode", ["Mock", "Real (Simulated OAuth)"], horizontal=True)

    if st.session_state.doc_type in supported_types:
        st.info(f"Document Type Supported: {st.session_state.doc_type.upper()}")

        if mode == "Real (Simulated OAuth)":
            if st.button("Authenticate with Concur"):
                st.session_state.concur_token = "mock_token"
                st.success("Authenticated")

        if st.button("Send to Concur"):
            progress = st.progress(0, text="Preparing payload...")

            payload = {
                "type": st.session_state.doc_type,
                "data": st.session_state.structured_data,
                "line_items": json_to_kv_dataframe(st.session_state.structured_data).to_dict(orient="records")
            }

            import time
            progress.progress(40, text="Connecting...")
            time.sleep(1)

            progress.progress(70, text="Sending...")
            time.sleep(1)

            if mode == "Mock":
                st.success("✅ Sent (Mock)")
            else:
                if "concur_token" not in st.session_state:
                    st.error("Authenticate first")
                    progress.empty()
                    st.stop()
                st.success("✅ Sent to API")

            progress.progress(100, text="Completed")
            progress.empty()

            st.json(payload)
    else:
        st.warning("Only Invoice or Ticket supported")

# METRICS

if selected_tab == "Metrics":
    st.subheader("📊 System Metrics")

    metrics = st.session_state.metrics

    if not st.session_state.processed_file:
        st.warning("Upload and process a document to view metrics.")

    total_calls = len(metrics["response_times"])
    avg_time = sum(metrics["response_times"]) / total_calls if total_calls else 0

    st.metric("Total Token Usage", metrics["tokens"])
    st.metric("Total LLM Calls", total_calls)
    st.metric("Avg Response Time (s)", round(avg_time, 2))

    # Basic accuracy placeholder (can improve later)
    if st.session_state.structured_data:
        accuracy = 0.85
        st.metric("Estimated Accuracy", f"{int(accuracy*100)}%")

    if metrics["response_times"]:
        df = pd.DataFrame({
            "Call #": list(range(len(metrics["response_times"]))),
            "Response Time (s)": metrics["response_times"]
        })

        st.line_chart(df.set_index("Call #"))

    else:
        st.info("No LLM calls yet. Upload a document to see metrics.")

    st.subheader("📂 Per Document Metrics")

    for file, data in st.session_state.doc_metrics.items():
        total_calls = data["calls"]
        avg_time = sum(data["response_times"]) / total_calls if total_calls else 0

        st.markdown(f"### 📄 {file}")
        st.write(f"Tokens: {data['tokens']}")
        st.write(f"Calls: {total_calls}")
        st.write(f"Avg Time: {round(avg_time, 2)} sec")

